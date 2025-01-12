import sys
from pathlib import Path
import modal
import io
import pandas as pd
from docx import Document as DocxDocument
from PIL import Image
import fitz
import os
import json

sys.path.append(str(Path(__file__).resolve().parents[2]))

from src.pipeline.registry import FunctionRegistry
from src.schemas.schemas import Document, Entry, Index, Ingestion, ParsedFeatureType, ParsingMethod, Scope, IngestionMethod
from src.utils.datetime_utils import get_current_utc_datetime

def convert_to_pdf(file_content: bytes, file_extension: str) -> bytes:
    pdf_bytes = io.BytesIO()
    if file_extension in ['.xlsx', '.xls']:
        # Convert Excel to PDF
        df = pd.read_excel(io.BytesIO(file_content))
        df.to_pdf(pdf_bytes)
    elif file_extension in ['.docx', '.doc']:
        # Convert Word to PDF
        doc = DocxDocument(io.BytesIO(file_content))
        # This is a placeholder. You'll need a library like python-docx2pdf for actual conversion
        # For now, we'll just extract text
        full_text = "\n".join([para.text for para in doc.paragraphs])
        pdf = fitz.open()
        page = pdf.new_page()
        page.insert_text((50, 50), full_text)
        pdf.save(pdf_bytes)
        pdf.close()
    elif file_extension in ['.png', '.jpg', '.jpeg']:
        # Convert Image to PDF
        image = Image.open(io.BytesIO(file_content))
        pdf = fitz.open()
        page = pdf.new_page(width=image.width, height=image.height)
        page.insert_image(page.rect, stream=file_content)
        pdf.save(pdf_bytes)
        pdf.close()
    else:
        # Assume it's already a PDF
        return file_content
    return pdf_bytes.getvalue()


@FunctionRegistry.register("parse", "datalab")
async def main_datalab(ingestions: list[Ingestion], write=None, read=None, mode="simple", **kwargs) -> list[Document]:
    file_bytes = []
    cls = modal.Cls.lookup("document-parsing-modal", "Model")
    obj = cls()
    for ingestion in ingestions:
        ingestion.parsing_method = ParsingMethod.MARKER
        ingestion.parsing_date = get_current_utc_datetime()
        ingestion.parsed_feature_type = [ParsedFeatureType.TEXT]

        # Set parsed_file_path (similar to ocr_service.py)
        if not ingestion.parsed_file_path:
            base_path = os.path.splitext(ingestion.file_path)[0]
            ingestion.parsed_file_path = f"{base_path}_datalab.txt"

        # Update file reading to handle async
        file_content = await read(ingestion.file_path, mode="rb") if read else open(ingestion.file_path, "rb").read()

        # Convert to PDF if necessary
        file_extension = Path(ingestion.file_path).suffix.lower()
        pdf_content = convert_to_pdf(file_content, file_extension)
        file_bytes.append(pdf_content)
        
        # Extract metadata
        with fitz.open(stream=io.BytesIO(pdf_content), filetype="pdf") as pdf:
            ingestion.metadata = pdf.metadata
    
    all_documents = []
    async for ret in obj.parse_document.map.aio(file_bytes, return_exceptions=True):
        document = Document(entries=[])
        if isinstance(ret, Exception):
            print(f"Error processing document: {ret}")
            continue

        parsed_data = json.loads(ret)  # Parse the JSON return from modal
        
        if mode == "simple":
            # Original behavior: concatenate all text on each page
            all_text = []
            for page_num, page in enumerate(parsed_data["children"]):
                page_text = []
                for child in page.get("children", []):
                    if "text" in child:
                        page_text.append(child["text"])
                
                combined_page_text = " ".join(page_text)
                page_entry = Entry(
                    ingestion=ingestion, 
                    string=combined_page_text, 
                    index_numbers=[Index(primary=page_num + 1)]
                )
                document.entries.append(page_entry)
                all_text.append(combined_page_text)
        
        else:  # mode == "structured"
            # New behavior: create entries based on semantic blocks
            for page_num, page in enumerate(parsed_data["children"]):
                for child in page.get("children", []):
                    block_type = child.get("block_type")
                    if "text" in child:
                        entry = Entry(
                            ingestion=ingestion,
                            string=child["text"],
                            index_numbers=[Index(primary=page_num + 1)],
                            metadata={
                                "block_type": block_type,
                                "polygon": child.get("polygon"),
                                "confidence": child.get("confidence")
                            }
                        )
                        document.entries.append(entry)
                        all_text.append(child["text"])

        # Save combined text to file
        combined_text = "\n\n=== PAGE BREAK ===\n\n".join(all_text)
        if write:
            await write(ingestion.parsed_file_path, combined_text, mode="w")
        else:
            with open(ingestion.parsed_file_path, "w", encoding='utf-8') as f:
                f.write(combined_text)

        all_documents.append(document)
    return all_documents
